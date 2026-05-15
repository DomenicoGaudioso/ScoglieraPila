# -*- coding: utf-8 -*-
from io import BytesIO
from typing import Iterable

import pandas as pd


def _to_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        return f"{value:.2f}"
    return str(value)


def _clean_df(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for col in out.columns:
        out[col] = out[col].map(_to_text)
    return out


def _set_styles(doc):
    from docx.shared import Pt, RGBColor

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    styles["Normal"].font.color.rgb = RGBColor(0, 0, 0)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(0, 0, 0)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(11)


def _add_table(doc, df: pd.DataFrame, title: str | None = None):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import RGBColor

    if title:
        doc.add_heading(title, level=2)
    df = _clean_df(df)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()


def _add_matplotlib_figure(doc, spec: dict):
    from docx.shared import Cm
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    df = spec.get("df")
    if df is None or df.empty:
        return
    x_col = spec.get("x")
    y_cols = spec.get("y")
    if isinstance(y_cols, str):
        y_cols = [y_cols]
    if not x_col or not y_cols:
        return

    title = spec.get("title", "Grafico")
    kind = spec.get("kind", "line")
    fig, ax = plt.subplots(figsize=(7.4, 3.9), dpi=160)
    try:
        x_values = df[x_col].astype(str) if kind == "bar" else pd.to_numeric(df[x_col], errors="coerce")
        if kind == "bar":
            y = pd.to_numeric(df[y_cols[0]], errors="coerce")
            ax.bar(x_values, y, color="#4b5563")
        else:
            for y_col in y_cols:
                y = pd.to_numeric(df[y_col], errors="coerce")
                ax.plot(x_values, y, marker="o", linewidth=1.8, label=str(y_col))
            if len(y_cols) > 1:
                ax.legend(frameon=False, fontsize=8)
        ax.set_title(title, loc="left", fontsize=10, fontweight="bold", color="#111111")
        ax.set_xlabel(str(x_col), color="#111111")
        ax.set_ylabel(spec.get("ylabel", "Valore"), color="#111111")
        ax.grid(True, color="#d1d5db", linewidth=0.6, alpha=0.8)
        ax.tick_params(axis="x", labelrotation=spec.get("rotation", 0), colors="#111111")
        ax.tick_params(axis="y", colors="#111111")
        for spine in ax.spines.values():
            spine.set_color("#9ca3af")
        fig.tight_layout()
        bio = BytesIO()
        fig.savefig(bio, format="png", dpi=160, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        bio.seek(0)
        doc.add_picture(bio, width=Cm(15.5))
        doc.add_paragraph(title)
    except Exception as exc:
        plt.close(fig)
        doc.add_paragraph(f"Figura non generata: {title}. Dettaglio tecnico: {exc}")

def _rows_to_df(rows: Iterable[dict]) -> pd.DataFrame:
    return pd.DataFrame(list(rows), columns=["Parametro", "Valore", "Unita", "Esito/nota"])


def genera_relazione_word(
    titolo: str,
    sottotitolo: str,
    input_rows: Iterable[dict],
    formule_df: pd.DataFrame | None,
    result_tables: list[tuple[str, pd.DataFrame]],
    note: Iterable[str] | None = None,
    figures: list[dict] | None = None,
) -> bytes:
    from docx import Document
    from docx.shared import Cm

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    _set_styles(doc)

    doc.add_heading(f"1. {titolo}", level=1)
    if sottotitolo:
        doc.add_paragraph(sottotitolo)

    doc.add_heading("1.1 Formulazione teorica", level=2)
    if formule_df is not None and not formule_df.empty:
        _add_table(doc, formule_df)
    else:
        doc.add_paragraph("Le formule operative sono riportate nelle tabelle di calcolo generate dalla app.")

    _add_table(doc, _rows_to_df(input_rows), "1.2 Dati di input")

    for idx, (title, df) in enumerate(result_tables, start=3):
        if df is not None and not df.empty:
            _add_table(doc, df, f"1.{idx} {title}")

    note_idx = len(result_tables) + 3
    if figures:
        doc.add_heading(f"1.{note_idx} Grafici statici", level=2)
        for spec in figures:
            _add_matplotlib_figure(doc, spec)
        note_idx += 1

    doc.add_heading(f"1.{note_idx} Note tecniche e limitazioni", level=2)
    if note:
        for item in note:
            doc.add_paragraph(str(item), style=None)
    else:
        doc.add_paragraph("Il documento riporta i risultati calcolati dalla app e deve essere verificato dal progettista.")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

